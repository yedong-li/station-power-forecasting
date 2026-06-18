import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


WORKSPACE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
SOURCE_MD = os.path.join(WORKSPACE_DIR, "基于Transformer与运维事件感知的电站短期功率预测方法.md")
OUTPUT_DOCX = os.path.join(WORKSPACE_DIR, "基于Transformer与运维事件感知的电站短期功率预测方法.docx")


def set_run_font(run, size=12, bold=False):
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def add_paragraph(doc, text, style=None, align=None, size=12, bold=False):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def clean_inline_markdown(text):
    text = text.strip()
    text = re.sub(r"<sup>(.*?)</sup>", r"\1", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("*", "")
    text = text.replace("`", "")
    return text.strip()


def parse_table(lines, start_idx):
    table_lines = []
    idx = start_idx
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1

    rows = []
    for line in table_lines:
        cells = [clean_inline_markdown(cell) for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=9, bold=(r_idx == 0))
    doc.add_paragraph()


def add_image(doc, image_path, caption=None):
    abs_path = os.path.join(WORKSPACE_DIR, image_path.replace("/", os.sep))
    if os.path.exists(abs_path):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        try:
            run.add_picture(abs_path, width=Inches(5.8))
        except Exception:
            add_paragraph(doc, f"[图片无法插入：{image_path}]", size=10)
    if caption:
        add_paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True)


def build_docx():
    with open(SOURCE_MD, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    idx = 0
    pending_image = None
    while idx < len(lines):
        raw = lines[idx].rstrip("\n")
        line = raw.strip()

        if not line or line == "---":
            idx += 1
            continue

        image_match = re.search(r'<img\s+src="([^"]+)"[^>]*>', line)
        if image_match:
            pending_image = image_match.group(1)
            idx += 1
            continue

        caption_match = re.search(r"<p><b>(.*?)</b></p>", line)
        if caption_match and pending_image:
            add_image(doc, pending_image, clean_inline_markdown(caption_match.group(1)))
            pending_image = None
            idx += 1
            continue

        if line.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue

        if line.startswith("# "):
            add_paragraph(doc, clean_inline_markdown(line[2:]), style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
        elif line.startswith("## "):
            add_paragraph(doc, clean_inline_markdown(line[3:]), style="Heading 1", size=14, bold=True)
        elif line.startswith("### "):
            add_paragraph(doc, clean_inline_markdown(line[4:]), style="Heading 2", size=12, bold=True)
        elif line.startswith("#### "):
            add_paragraph(doc, clean_inline_markdown(line[5:]), style="Heading 3", size=11, bold=True)
        elif line.startswith("##### "):
            add_paragraph(doc, clean_inline_markdown(line[6:]), align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True)
        elif line.startswith(">"):
            add_paragraph(doc, clean_inline_markdown(line.lstrip("> ")), size=11)
        elif line.startswith("<div") or line.startswith("</div>"):
            pass
        else:
            add_paragraph(doc, clean_inline_markdown(line), size=12)
        idx += 1

    doc.save(OUTPUT_DOCX)
    print(f"Generated DOCX: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_docx()
